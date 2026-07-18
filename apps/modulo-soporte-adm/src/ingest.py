"""
Pipeline de ingesta:
  1. Carga documentos del corpus (.md, .txt, .docx)
  2. Divide el texto en chunks con solapamiento (respetando encabezados)
  3. Genera embeddings y los almacena en ChromaDB (persistente)

Uso:
    python -m src.ingest              # ingesta todo data/corpus/
    python -m src.ingest --reset      # borra la colección y re-ingesta
"""
import argparse
import hashlib
import re
from pathlib import Path
from typing import Dict, List, Tuple

import chromadb
from docx import Document as DocxDocument

from .config import (
    CHROMA_DIR,
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    COLLECTION_NAME,
    CORPUS_DIR,
)
from .embeddings import get_embedding_function

HEADER_RE = re.compile(r"^(#{1,6})\s+(.*)$")


# ---------------------------------------------------------------- carga
def load_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading"):
            level = "".join(filter(str.isdigit, p.style.name)) or "2"
            parts.append(f"{'#' * int(level)} {text}")
        else:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def load_document(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return load_docx(path)
    return path.read_text(encoding="utf-8", errors="ignore")


# ---------------------------------------------------------------- chunking
def split_by_headers(text: str) -> List[Tuple[str, str]]:
    """Divide el texto en secciones (encabezado_actual, contenido)."""
    sections: List[Tuple[str, str]] = []
    current_header = "INICIO"
    buffer: List[str] = []
    for line in text.splitlines():
        m = HEADER_RE.match(line.strip())
        if m:
            if buffer:
                sections.append((current_header, "\n".join(buffer).strip()))
                buffer = []
            current_header = m.group(2).strip("* \t")
        else:
            buffer.append(line)
    if buffer:
        sections.append((current_header, "\n".join(buffer).strip()))
    return [(h, c) for h, c in sections if c]


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Chunking recursivo simple por párrafos con solapamiento."""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: List[str] = []
    current = ""
    for p in paragraphs:
        if len(current) + len(p) + 2 <= size:
            current = f"{current}\n\n{p}" if current else p
        else:
            if current:
                chunks.append(current)
            # párrafo más grande que el chunk: cortar duro con solapamiento
            while len(p) > size:
                chunks.append(p[:size])
                p = p[size - overlap:]
            current = p
    if current:
        chunks.append(current)

    # aplicar solapamiento entre chunks consecutivos
    overlapped: List[str] = []
    for i, ch in enumerate(chunks):
        if i > 0 and overlap > 0:
            tail = chunks[i - 1][-overlap:]
            ch = f"{tail}\n{ch}"
        overlapped.append(ch)
    return overlapped


def build_chunks(path: Path) -> Tuple[List[str], List[Dict], List[str]]:
    """Devuelve (documentos, metadatos, ids) listos para ChromaDB."""
    text = load_document(path)
    docs: List[str] = []
    metas: List[Dict] = []
    ids: List[str] = []
    for header, content in split_by_headers(text):
        for j, chunk in enumerate(chunk_text(content)):
            enriched = f"[Sección: {header}]\n{chunk}"
            uid = hashlib.md5(f"{path.name}|{header}|{j}|{chunk[:64]}".encode()).hexdigest()
            docs.append(enriched)
            metas.append({"source": path.name, "section": header, "chunk_index": j})
            ids.append(uid)
    return docs, metas, ids


# ---------------------------------------------------------------- ingesta
def get_collection(reset: bool = False):
    client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    if reset:
        try:
            client.delete_collection(COLLECTION_NAME)
        except Exception:
            pass
    return client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=get_embedding_function(),
        metadata={"hnsw:space": "cosine"},
    )


def ingest_corpus(reset: bool = False) -> int:
    collection = get_collection(reset=reset)
    total = 0
    files = sorted(
        [p for p in CORPUS_DIR.glob("**/*") if p.suffix.lower() in {".md", ".txt", ".docx"}]
    )
    if not files:
        print(f"[!] No se encontraron documentos en {CORPUS_DIR}")
        return 0
    for path in files:
        docs, metas, ids = build_chunks(path)
        if not docs:
            continue
        BATCH = 64
        for i in range(0, len(docs), BATCH):
            collection.upsert(
                documents=docs[i : i + BATCH],
                metadatas=metas[i : i + BATCH],
                ids=ids[i : i + BATCH],
            )
        print(f"[+] {path.name}: {len(docs)} chunks ingresados")
        total += len(docs)
    print(f"[✓] Ingesta completa: {total} chunks | colección '{COLLECTION_NAME}'")
    return total


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ingesta del corpus al vector store")
    parser.add_argument("--reset", action="store_true", help="Recrear la colección desde cero")
    args = parser.parse_args()
    ingest_corpus(reset=args.reset)
