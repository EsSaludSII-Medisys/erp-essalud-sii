"""
generar_docx.py — Genera INFORME.docx (Word) a partir del contenido del informe.

Crea un documento Word con títulos, tablas, bloques de código y espacios
marcados para pegar las capturas de pantalla.

Requisitos:
    pip install python-docx

Uso:
    python generar_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Estilos base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

AZUL = RGBColor(0x16, 0x32, 0x4F)


def h1(txt):
    p = doc.add_heading(txt, level=0)
    return p


def h2(txt):
    doc.add_heading(txt, level=1)


def h3(txt):
    doc.add_heading(txt, level=2)


def parrafo(txt):
    doc.add_paragraph(txt)


def codigo(txt):
    p = doc.add_paragraph()
    run = p.add_run(txt)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    # sombreado gris claro
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)


def tabla(headers, filas):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for fila in filas:
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = val


def placeholder_imagen(texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {texto} ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)


# ======================= CONTENIDO =======================

h1("Informe: Búsqueda de Información en Documentos mediante RAG")

parrafo("Trabajo individual")
parrafo("Autor: Diego Jesús Medina Tirado")
parrafo("Universidad Privada Antenor Orrego (UPAO)")
parrafo("Fecha: 3 de julio de 2026")
p = doc.add_paragraph()
p.add_run("Repositorio con la solución: ").bold = True
p.add_run("https://github.com/Medina-Tirado-Diego/rag-sii-upao")

h2("1. Planteamiento del problema")
parrafo(
    "La universidad requiere un programa que permita realizar búsqueda de "
    "información sobre documentos alojados en un repositorio. El usuario debe "
    "poder ingresar una pregunta en lenguaje natural, por ejemplo: "
    "\"Requisitos para reserva de matrícula\", y el sistema debe devolver "
    "información pertinente sobre ese tema, extraída de los documentos oficiales "
    "(reglamentos, guías de trámites, etc.). Como propuesta tecnológica para el "
    "futuro producto software se emplea la técnica de IA denominada RAG "
    "(Retrieval-Augmented Generation)."
)

h2("2. ¿Qué es RAG? (Retrieval-Augmented Generation)")
parrafo(
    "RAG combina un motor de búsqueda semántica con un modelo de lenguaje (LLM). "
    "En lugar de que el modelo responda solo con su conocimiento interno (lo que "
    "puede producir invenciones o 'alucinaciones'), primero se recuperan los "
    "fragmentos de documento más relevantes y luego se le pide al LLM que genere "
    "la respuesta usando esa información como contexto."
)
h3("Fases")
parrafo(
    "1. Recuperación (Retrieval): la pregunta se transforma en un embedding "
    "(vector numérico que representa su significado). Se compara con los vectores "
    "de todos los fragmentos almacenados en una base de datos vectorial y se "
    "seleccionan los más similares."
)
parrafo(
    "2. Generación (Generation): los fragmentos recuperados se insertan como "
    "contexto dentro del prompt que recibe el LLM, el cual redacta la respuesta "
    "final basándose únicamente en dicha información."
)
h3("Arquitectura")
codigo(
    "FASE DE INDEXACIÓN (una sola vez):\n"
    "Documentos -> Fragmentar (chunks) -> Embeddings -> Vector DB\n\n"
    "FASE DE CONSULTA (cada pregunta):\n"
    "Pregunta -> Embedding -> Búsqueda por similitud -> Chunks relevantes\n"
    "                        Chunks + Pregunta -> LLM -> Respuesta final"
)
h3("¿Por qué RAG? Ventajas para el producto futuro")
tabla(
    ["Ventaja", "Beneficio para la universidad"],
    [
        ["Respuestas basadas en fuentes reales", "Reduce las alucinaciones; la información sale de los reglamentos oficiales."],
        ["Actualizable sin reentrenar", "Basta con añadir/editar documentos y re-indexar; sin reentrenamiento costoso."],
        ["Trazabilidad", "El sistema indica de qué documento proviene la respuesta."],
        ["Escalabilidad", "De 3 documentos de prueba se puede crecer a todo el reglamento institucional."],
        ["Bajo costo / privacidad", "Corre 100% local, sin enviar datos a terceros ni pagar APIs."],
    ],
)

h2("3. Arquitectura y stack tecnológico implementado")
parrafo("Solución 100% local y gratuita (sin claves de API):")
tabla(
    ["Componente", "Tecnología", "Rol"],
    [
        ["Orquestación del flujo RAG", "LangChain", "Conecta carga, fragmentación, embeddings, búsqueda y LLM."],
        ["Base de datos vectorial", "ChromaDB", "Almacena embeddings y busca por similitud (local, sin servidor)."],
        ["Modelo de embeddings", "sentence-transformers (paraphrase-multilingual-MiniLM-L12-v2)", "Convierte texto en vectores; multilingüe (español)."],
        ["Modelo de lenguaje (LLM)", "Ollama con llama3.2", "Genera la respuesta final localmente."],
        ["Lenguaje", "Python 3.13", "Implementación de los scripts."],
    ],
)

h2("4. Estructura del proyecto y archivos .PY")
codigo(
    "rag-sii-upao/\n"
    "├── documentos/            # Repositorio de documentos (.txt)\n"
    "│   ├── matricula.txt       #   Reglamento de matrícula (reserva de matrícula)\n"
    "│   ├── becas.txt           #   Reglamento de becas\n"
    "│   └── tramites.txt        #   Guía de trámites del SII\n"
    "├── ingest.py              # (.PY) Indexa los documentos en ChromaDB\n"
    "├── query.py               # (.PY) Recibe la pregunta y genera la respuesta\n"
    "├── requirements.txt       # Paquetes usados con versión\n"
    "└── README.md              # Documentación de instalación y uso"
)
parrafo(
    "Los dos archivos .py que evidencian la propuesta son ingest.py (indexación) "
    "y query.py (consulta). Su código completo se anexa en la sección 8."
)

h2("5. Prompts usados como apoyo en la generación del RAG")
h3("5.1. Prompt dado a la IA (Claude) para generar el código")
codigo(
    "Construir un sistema RAG en Python donde hay documentos de la universidad\n"
    "(reglamentos, requisitos de matrícula, procedimientos). El usuario hace una\n"
    "pregunta en lenguaje natural (ej: '¿Cuáles son los requisitos para reserva de\n"
    "matrícula?') y el sistema busca en los documentos y devuelve una respuesta\n"
    "basada en esa información. Usar LangChain + ChromaDB. Embeddings locales con\n"
    "sentence-transformers y LLM local con Ollama (sin API keys). Crear ingest.py\n"
    "que carga los documentos, los fragmenta en chunks de 500 caracteres con\n"
    "solapamiento y los guarda en ChromaDB; y query.py que recupera los chunks más\n"
    "similares, arma un prompt con ellos como contexto y se lo pasa al LLM para la\n"
    "respuesta final. Preparar 2-3 documentos de prueba (matrícula, becas, trámites)."
)
h3("5.2. Prompt template interno del RAG (dentro de query.py)")
parrafo(
    "Prompt que el sistema construye automáticamente y envía al LLM en cada "
    "consulta. Obliga al modelo a responder solo con el contexto recuperado:"
)
codigo(
    "SYSTEM:\n"
    "Eres un asistente académico de la Universidad Privada Antenor Orrego (UPAO).\n"
    "Responde la pregunta del estudiante ÚNICAMENTE con la información del CONTEXTO\n"
    "proporcionado. Si la respuesta no está en el contexto, di claramente que no\n"
    "cuentas con esa información. Responde en español, de forma clara y, cuando\n"
    "corresponda, en forma de lista.\n\n"
    "HUMAN:\n"
    "CONTEXTO:\n{contexto}\n\n"
    "PREGUNTA: {pregunta}\n\n"
    "RESPUESTA:"
)

h2("6. Paquetes usados (dependencias)")
tabla(
    ["Paquete", "Versión"],
    [
        ["langchain", "1.3.11"],
        ["langchain-community", "0.4.2"],
        ["langchain-text-splitters", "1.1.2"],
        ["langchain-huggingface", "1.2.2"],
        ["langchain-chroma", "1.1.0"],
        ["langchain-ollama", "1.1.0"],
        ["chromadb", "1.5.9"],
        ["sentence-transformers", "5.6.0"],
        ["transformers", "5.13.0"],
        ["torch", "2.12.1+cpu"],
        ["ollama", "0.6.2"],
    ],
)
parrafo("La lista completa y exacta está en el archivo requirements.txt (pip freeze).")
h3("Casos/problemas encontrados y su solución")
tabla(
    ["Caso", "Solución aplicada"],
    [
        ["ModuleNotFoundError al ejecutar los scripts", "El intérprete python no apuntaba al entorno virtual; se recreó el .venv y se instalaron las dependencias en él."],
        ["Error [Errno 28] No space left on device", "torch arrastra ~2.5 GB de CUDA que llenaban /tmp. Se instaló torch CPU (~192 MB) desde download.pytorch.org/whl/cpu."],
        ["/tmp demasiado pequeño en la instalación", "Se redirigió el temporal al disco principal con TMPDIR=~/tmp-pip."],
        ["Ollama no instalado", "Se instaló como servicio del sistema y se descargó el modelo con 'ollama pull llama3.2'."],
    ],
)

h2("7. Ejecución y evidencia de funcionamiento")
parrafo("Comandos:")
codigo(
    'python ingest.py\n'
    'python query.py "¿Cuáles son los requisitos para reserva de matrícula?"'
)
parrafo("Salida real del programa (caso de ejemplo del enunciado):")
codigo(
    "=== PREGUNTA: ¿Cuáles son los requisitos para reserva de matrícula? ===\n\n"
    "--- RESPUESTA ---\n"
    "Según el Fragmento 3 y Fragmento 4, los requisitos para la reserva de\n"
    "matrícula son:\n"
    "* Solicitud dirigida al Director de la Escuela Profesional a través del SII.\n"
    "* No tener deudas pendientes con la universidad (pensiones, biblioteca, lab.).\n"
    "* Haberse matriculado al menos en un semestre académico anterior.\n"
    "* Pago de la tasa administrativa por reserva de matrícula vigente.\n"
    "* La reserva se concede por un máximo de dos (2) semestres académicos.\n\n"
    "--- FUENTES CONSULTADAS ---\n"
    "  * documentos/matricula.txt (x4)"
)
parrafo("Captura de pantalla de la ejecución:")
placeholder_imagen("PEGA AQUÍ LA CAPTURA DE PANTALLA DE LA TERMINAL")
parrafo(
    "El sistema recuperó correctamente los fragmentos del documento matricula.txt "
    "y el LLM generó una respuesta fiel a la fuente, sin inventar información."
)

h2("8. Anexo: código fuente de los archivos .PY")
h3("8.1. ingest.py — Indexación")
codigo(open("ingest.py", encoding="utf-8").read())
h3("8.2. query.py — Consulta")
codigo(open("query.py", encoding="utf-8").read())

h2("9. Conclusión")
parrafo(
    "Se implementó y probó con éxito un sistema RAG que responde preguntas en "
    "lenguaje natural sobre documentos de la universidad, funcionando de forma "
    "local y gratuita. El caso de ejemplo ('Requisitos para reserva de matrícula') "
    "se resolvió correctamente, recuperando la información del reglamento y "
    "generando una respuesta fiel a la fuente. Esta arquitectura constituye una "
    "base sólida y escalable para el futuro producto software de la universidad."
)

doc.save("INFORME.docx")
print("Listo: INFORME.docx")
